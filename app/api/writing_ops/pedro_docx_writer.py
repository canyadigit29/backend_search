
from docx import Document
from docx.shared import Inches
import os

def generate_docx(descriptions, chart_paths, output_path):
    doc = Document()
    doc.add_heading("Generated Report", level=1)

    # Executive Summary
    summary = descriptions.get("executive_summary", "")
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(summary)

    # Sectioned Analysis
    for section in descriptions.get("sections", []):
        title = section.get("title", "Untitled Section")
        text = section.get("content", "")
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

    # Insert charts
    for chart in chart_paths:
        if os.path.exists(chart):
            doc.add_picture(chart, width=Inches(5.5))

    doc.save(output_path)
